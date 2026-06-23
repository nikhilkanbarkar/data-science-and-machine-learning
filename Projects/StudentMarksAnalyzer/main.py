#pip install faker      #if already installed then comment it!

import faker as Faker
from faker import Faker
import numpy as np
import pandas as pd
from docx import Document


# Student is Pass or Fail
def check_result(row):
    failed_subjects=[]

    for subject in subjects:

        if row[subject] < 35:
            failed_subjects.append(subject)
    if len(failed_subjects) == 0:
        return pd.Series(["Pass","None"])

    else:
        return pd.Series([
            "Fail",
            ", ".join(failed_subjects)
        ])


# Grade Calculation
def check_grade(row):
    if row["Result"] == "Fail":
        return "F"
    else:
        if row["Percentage"] >= 90:
            return "A+"
        elif row["Percentage"] >= 80:
            return "A"
        elif row["Percentage"] >= 70:
            return "B+"
        elif row["Percentage"] >= 60:
            return "B"
        elif row["Percentage"] >= 50:
            return "C+"
        else:
            return "C"


# Topppers 10 of the class
def toppers_of_class():
    passed_students = dataframe[
        dataframe["Result"] == "Pass"
    ]
    top_10 = passed_students.sort_values(
        by = "Total_marks",
        ascending = False
    ).head(10)
    return top_10


# Subject-wise Analysis
def subject_analysis(dataframe, subject):

    total_students = len(dataframe)

    topper = dataframe.loc[
        dataframe[subject].idxmax()
    ]

    passed = (
        dataframe[subject] >= 35
    ).sum()

    failed = (
        dataframe[subject] < 35
    ).sum()

    pass_percentage = (
        passed / total_students
    ) * 100

    fail_percentage = (
        failed / total_students
    ) * 100

    average_marks = dataframe[
        subject
    ].mean()

    highest_marks = dataframe[
        subject
    ].max()

    lowest_marks = dataframe[
        subject
    ].min()

    if pass_percentage >= 85:
        status = "Good"

    elif pass_percentage >= 70:
        status = "Average"

    else:
        status = "Poor"
    return {
    "Subject": subject,
    "Topper": topper["Names"],
    "Topper Marks": highest_marks,
    "Passed": passed,
    "Failed": failed,
    "Pass %": pass_percentage,
    "Fail %": fail_percentage,
    "Average": average_marks,
    "Status": status
    }

def save_report_docx():

    document = Document()

    document.add_heading(
        'Student Performance Analysis Report',
        level=1
    )

    document.add_heading(
        'Overall Statistics',
        level=2
    )

    document.add_paragraph(
        f"Total Students: {len(dataframe)}"
    )

    document.add_paragraph(
        f"Passed Students: {No_of_students_passed}"
    )

    document.add_paragraph(
        f"Failed Students: {No_of_students_failed}"
    )

    document.add_paragraph(
        f"Pass Percentage: {Passed_percentage:.2f}%"
    )

    document.add_paragraph(
        f"Fail Percentage: {Failed_percentage:.2f}%"
    )


    topper = top_10.iloc[0]

    document.add_heading(
        'Overall Topper',
        level=2
    )

    document.add_paragraph(
        f"Name: {topper['Names']}"
    )

    document.add_paragraph(
        f"Roll No: {topper['Roll_No']}"
    )

    document.add_paragraph(
        f"Total Marks: {topper['Total_marks']}"
    )

    document.add_paragraph(
        f"Percentage: {topper['Percentage']:.2f}%"
    )

    document.add_paragraph(
        f"Grade: {topper['Grade']}"
    )

    document.add_heading(
        'Top 10 Toppers',
        level=2
    )

    document.add_paragraph(
        top_10[
            [
                "Roll_No",
                "Names",
                "Total_marks",
                "Percentage",
                "Grade"
            ]
        ].to_string()
    )


    document.add_heading(
        'Subject Wise Analysis',
        level=2
    )

    document.add_paragraph(
        subjectwise_report_df.to_string()
    )

    easiest_subject = subjectwise_report_df.loc[
        subjectwise_report_df["Pass %"].idxmax(),
        "Subject"
    ]

    difficult_subject = subjectwise_report_df.loc[
        subjectwise_report_df["Pass %"].idxmin(),
        "Subject"
    ]

    document.add_heading(
        'Key Insights',
        level=2
    )

    document.add_paragraph(
        f"Easiest Subject: {easiest_subject}"
    )

    document.add_paragraph(
        f"Most Difficult Subject: {difficult_subject}"
    )

    document.save(
        "Student_Performance_Report.docx"
    )

    print(
        "Report Saved Successfully!"
    )

def print_final_report():

    print("\n")
    print("=" * 70)
    print("STUDENT PERFORMANCE ANALYSIS REPORT")
    print("=" * 70)

    print(f"Total Students           : {len(dataframe)}")
    print(f"Passed Students          : {No_of_students_passed}")
    print(f"Failed Students          : {No_of_students_failed}")

    print(f"Pass Percentage          : {Passed_percentage:.2f}%")
    print(f"Fail Percentage          : {Failed_percentage:.2f}%")

    topper = top_10.iloc[0]

    print("\n")
    print("=" * 70)
    print("OVERALL TOPPER")
    print("=" * 70)

    print(f"Name                     : {topper['Names']}")
    print(f"Roll No                  : {topper['Roll_No']}")
    print(f"Total Marks              : {topper['Total_marks']}")
    print(f"Percentage               : {topper['Percentage']:.2f}%")
    print(f"Grade                    : {topper['Grade']}")

    print("\n")
    print("=" * 70)
    print("TOP 10 TOPPERS")
    print("=" * 70)

    print(
        top_10[
            [
                "Roll_No",
                "Names",
                "Total_marks",
                "Percentage",
                "Grade"
            ]
        ].to_string(index=False)
    )

    print("\n")
    print("=" * 70)
    print("SUBJECT WISE ANALYSIS")
    print("=" * 70)

    print(
        subjectwise_report_df.to_string(index=False)
    )

    easiest_subject = subjectwise_report_df.loc[
        subjectwise_report_df["Pass %"].idxmax(),
        "Subject"
    ]

    difficult_subject = subjectwise_report_df.loc[
        subjectwise_report_df["Pass %"].idxmin(),
        "Subject"
    ]

    print("\n")
    print("=" * 70)
    print("KEY INSIGHTS")
    print("=" * 70)

    print(
        f"Easiest Subject          : {easiest_subject}"
    )

    print(
        f"Most Difficult Subject   : {difficult_subject}"
    )

    print("\n")
    print("=" * 70)
    print("REPORT GENERATED SUCCESSFULLY")
    print("=" * 70)


'''******************************************************************************************************************************************************************'''
# main code
fake = Faker()
names = [fake.first_name() for _ in range(500)]

#Sort the names A-->Z
names.sort()

# Now create roll_numbers for the students
roll_numbers = np.arange(2026001, 2026001 + 500)

# Let's create a Matrix of rando numbers in order to create marks of the students
np.random.seed(0)       
''' Above line will keep the Same 
randomised numbers which u have generated at 1st time '''
marks = np.random.randint(25,101,size = (500,6)) # 500*6 Matrix


# Let's create a DataFrame
dataframe = pd.DataFrame(marks,
                        #rows = names,
                        columns = ["Maths","Science","Social_sci","Marathi","Kannada","English"])

# Lets add Names and the Roll_No column in the DataFrame
dataframe["Names"] = names
dataframe["Roll_No"] = roll_numbers

# Below lines will re-organize the DataFramecolumns
dataframe = dataframe[[
    "Roll_No",
    "Names",
    "Maths",
    "Science",
    "Social_sci",
    "English",
    "Marathi",
    "Kannada"
]]


# Calculating the total marks 
subjects =[
    "Maths",
    "Science",
    "Social_sci",
    "English",
    "Marathi",
    "Kannada"
]
dataframe["Total_marks"] = dataframe[subjects].sum(axis=1)


# Calculating the % of the students
dataframe["Percentage"] = (dataframe["Total_marks"]/600)*100

# Pass or Fail status and failed subjects
dataframe[["Result", "Failed_subjects"]] = dataframe.apply(
        check_result,
        axis = 1
    )

# Grade Calculation
dataframe["Grade"] = dataframe.apply(
    check_grade,
    axis = 1
)


# passed and failed students DataFrame and the percentage
Passed_students = dataframe[
        dataframe["Result"] == "Pass"
    ]
Failed_students = dataframe[
        dataframe["Result"] == "Fail"
    ]
No_of_students_passed = len(Passed_students)
No_of_students_failed = len(Failed_students)

Passed_percentage = No_of_students_passed/500*100
Failed_percentage = No_of_students_failed/500*100


# Top 10 list
top_10 = toppers_of_class()


# Subject-wise Analysis 
analysis = []

for subject in subjects:
    analysis.append(
        subject_analysis(dataframe, subject)
    )

subjectwise_report_df = pd.DataFrame(analysis, index = [1,2,3,4,5,6])


# complete report is as follows

# Save final dataframe as .csv file
dataframe.to_csv("Class_Report.csv")

# Passed students
Passed_students.to_csv("Passed_students.csv")

# Failed students
Failed_students.to_csv("Failed_students.csv")

# Toppers list from the class
top_10.to_csv("Top_10_list.csv")

# Subject-wise Analysis
subjectwise_report_df.to_csv("Subject_wise_Analysis.csv")

# Final report printing and saving
print_final_report()
save_report_docx()

